import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import time
 
#function for extracting emails from word document
def emails():
    try:
        doc_name = input("What is the name of the docx file containing the emails ").strip()
        if '.docx' not in doc_name:
            doc_name += '.docx'

        email_lst = []
        doc = docx.Document(doc_name) 
        rels = doc.part.rels       
        for rel in rels:
            if rels[rel].reltype == RT.HYPERLINK:
                email = rels[rel]._target[7:]
                email_lst.append(email)
        return email_lst
    except docx.opc.exceptions.PackageNotFoundError:
        print('The file name you typed does not exist. Please ensure you typed the right name')
        time.sleep(2)
        print(emails())
    except Exception as e:
        print("An error occured. Try closing the program and running it again. This is the error: ", e)
        time.sleep(2)
        print(emails())

#function for extracting text from word document
def content():
    try:
        doc_name = input("What is the name of the docx file containing the content of the mail ").strip()
        if '.docx' not in doc_name:
            doc_name += '.docx'

        doc = docx.Document(doc_name)
        
        subject = doc.paragraphs[0].text
        body = ''
        for paragraph in doc.paragraphs[1:]:
            body += paragraph.text
            body += '\n'
        return subject, body

    except docx.opc.exceptions.PackageNotFoundError:
        print('The file name you typed does not exist. Please ensure you typed the right name')
        print(content())
    except Exception as e:
        print("An error occured. Try closing the program and running it again. This is the error: ", e)
        print(content())

